#!/usr/bin/python3
# -*- coding: utf-8 -*-

import sys
import functools
import os    
from docx import Document
from docx.shared import Inches

version = "0.1"

document = Document()
open_file = functools.partial(open, encoding='utf-8')
for filename in sorted(os.listdir(os.getcwd())):
    if filename.endswith(".jpg"):
        # print("Currently working on "+filename+"\n")
        document.add_picture(filename, width=Inches(5))
        document.add_page_break()
    elif filename.endswith(".txt"):
        # print ("Currently working on "+filename+"\n")
        f = open_file(filename, "r")
        document.add_paragraph(f.read())
        f.close()
        document.add_page_break()
document.save('OCRed_Book.docx')
